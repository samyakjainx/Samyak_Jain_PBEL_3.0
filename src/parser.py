"""
parser.py
----------
Handles extraction of raw text from uploaded resume files.

Supports:
    - PDF files (using PyMuPDF as primary engine, pdfplumber as fallback)
    - DOCX files (using python-docx)

Design notes:
    - Every function returns a plain string of extracted text.
    - Functions never raise on a single bad file; they return an
      empty string and log the issue, so one broken resume never
      crashes the whole batch-screening process.
"""

import io
import docx
import fitz  # PyMuPDF
import pdfplumber


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using PyMuPDF first.
    Falls back to pdfplumber if PyMuPDF returns no usable text
    (this can happen with certain PDF encodings/layouts).

    Args:
        file_bytes: raw bytes of the uploaded PDF file.

    Returns:
        Extracted text as a single string (empty string on failure).
    """
    text = ""

    # --- Primary engine: PyMuPDF (fast, works for most PDFs) ---
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            pages_text = [page.get_text("text") for page in doc]
            text = "\n".join(pages_text).strip()
    except Exception:
        text = ""

    # --- Fallback engine: pdfplumber (better for tricky layouts) ---
    if not text:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages_text = [page.extract_text() or "" for page in pdf.pages]
                text = "\n".join(pages_text).strip()
        except Exception:
            text = ""

    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file_bytes: raw bytes of the uploaded DOCX file.

    Returns:
        Extracted text as a single string (empty string on failure).
    """
    try:
        document = docx.Document(io.BytesIO(file_bytes))

        # Paragraph text
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        # Table text (resumes sometimes use tables for layout)
        table_text = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_text.append(cell.text.strip())

        full_text = "\n".join(paragraphs + table_text)
        return full_text.strip()
    except Exception:
        return ""


def extract_text(uploaded_file) -> str:
    """
    Universal extractor that detects file type from the filename
    extension and routes to the correct parser.

    Args:
        uploaded_file: a Streamlit UploadedFile object (has .name and
                       .read() / getvalue()).

    Returns:
        Extracted text as a string. Returns an empty string if the
        file type is unsupported or parsing fails.
    """
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        return ""


def is_supported_file(filename: str) -> bool:
    """Quick helper to validate file extension before processing."""
    filename = filename.lower()
    return filename.endswith(".pdf") or filename.endswith(".docx")
